from PyQt5.QtWidgets import QMainWindow, QApplication, QDialog
from PyQt5.QtWidgets import QTableWidgetItem
from backup_db import *
from all_models import *
from adm_db_panel import Ui_MainWindow
from login_panel import login_panell
from pols_menu import pols_main_menu
from admin_menu import adm_menu
from create_pers import create_per
from descr_spell import descript_spell
from descr import descrip
import random
import docx
import pandas as pd


# Панель администратора базы данных
class Admin_db_Panel(QMainWindow):
    def __init__(self, parent=None):
        super(Admin_db_Panel, self).__init__(parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.show_table())
        self.ui.pushButton_2.clicked.connect(lambda: self.change_table())
        self.ui.pushButton_3.clicked.connect(lambda: self.confirm_change())
        self.ui.pushButton_4.clicked.connect(lambda: self.undo_change())
        self.ui.action.triggered.connect(lambda: self.exit_db_panel())

    # Вывод таблицы
    def show_table(self):
        self.ui.comboBox_2.clear()
        self.ui.comboBox_3.clear()
        self.ui.listView.clearSelection()
        result = self.ui.comboBox.currentText()
        table = Weapon
        line = 0
        collums = []
        if result == "Типы оружия":
            table = weapon_types
            line = session.query(weapon_types).count()
            self.ui.listView.setColumnCount(2)
            self.ui.qTable = session.query(weapon_types).all()
            collums = ['id', 'name']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Название"])
        if result == "Типы брони":
            table = armor_types
            line = session.query(armor_types).count()
            self.ui.listView.setColumnCount(2)
            self.ui.qTable = session.query(armor_types).all()
            collums = ['id', 'name']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Название"])
        if result == "Статусы пользователя":
            table = user_status
            line = session.query(user_status).count()
            self.ui.listView.setColumnCount(2)
            self.ui.qTable = session.query(user_status).all()
            collums = ['id', 'status']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Статус"])
        if result == "Описания":
            table = descriptions
            line = session.query(descriptions).count()
            self.ui.listView.setColumnCount(2)
            self.ui.qTable = session.query(descriptions).all()
            collums = ['id', 'field']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Описание"])
        if result == "Расы":
            table = Races
            line = session.query(Races).count()
            self.ui.listView.setColumnCount(7)
            self.ui.qTable = session.query(Races).all()
            collums = ['id', 'name', 'incr_char', 'worldview', 'size', 'speed', 'descr_id']
            self.ui.listView.setHorizontalHeaderLabels(
                ["ID", "Название", "Повышение характеристик", "Мировозрение", "Размер", "Скорость", "ID описания"])
        if result == "Разновидности рас":
            table = var_races
            line = 14  # session.query(var_races).count()
            self.ui.listView.setColumnCount(6)
            self.ui.qTable = session.query(var_races).all()
            collums = ['id', 'name', 'incr_char', 'add_feat', 'rac_id', 'descr_id']
            self.ui.listView.setHorizontalHeaderLabels(
                ["ID", "Название", "Повышение характеристик", "Доп. способность", "ID расы", "ID описания"])
        if result == "Аккаунты":
            table = Accounts
            line = session.query(Accounts).count()
            self.ui.listView.setColumnCount(4)
            self.ui.qTable = session.query(Accounts).all()
            collums = ['id', 'login', 'password', 'stat_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Логин", "Пароль", "ID статуса"])
        if result == "Классы":
            table = Classes
            line = session.query(Classes).count()
            self.ui.listView.setColumnCount(5)
            self.ui.qTable = session.query(Classes).all()
            collums = ['id', 'name', 'master_bonus', 'numb_av_spells', 'descr_id']
            self.ui.listView.setHorizontalHeaderLabels(
                ["ID", "Название", "Бонус мастерства", "Кол-во доступных заклинаний", "ID описания"])
        if result == "Оружие":
            table = Weapon
            line = session.query(Weapon).count()
            self.ui.listView.setColumnCount(6)
            self.ui.qTable = session.query(Weapon).all()
            collums = ['id', 'name', 'price', 'damage', 'weight', 'weap_t_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Название", "Цена", "Урон", "Вес", "ID типа"])
        if result == "Броня":
            table = Armor
            line = session.query(Armor).count()
            self.ui.listView.setColumnCount(6)
            self.ui.qTable = session.query(Armor).all()
            collums = ['id', 'name', 'price', 'steal_hindr', 'weight', 'arm_t_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "Название", "Цена", "Громкий", "Вес", "ID типа"])
        if result == "Таблица заклинаний":
            table = Spell_table
            line = session.query(Spell_table).count()
            self.ui.listView.setColumnCount(6)
            self.ui.qTable = session.query(Spell_table).all()
            collums = ['id', 'name', 'level', 'distance', 'duration', 'descr_id']
            self.ui.listView.setHorizontalHeaderLabels(
                ["ID", "Название", "Уровень", "Дистанция", "Длительность", "ID описания"])
        if result == "Связи типов оружия и классов":
            table = Relat_table_t_weap_t_cl
            line = 21  # session.query(Relat_table_t_weap_t_cl).count()
            self.ui.listView.setColumnCount(3)
            self.ui.qTable = session.query(Relat_table_t_weap_t_cl).all()
            collums = ['id', 'weap_t_id', 'class_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "ID типа оружия", "ID класса"])
        if result == "Связи типов доспехов и классов":
            table = Relat_table_t_arm_t_cl
            line = 25  # session.query(Relat_table_t_arm_t_cl).count()
            self.ui.listView.setColumnCount(3)
            self.ui.qTable = session.query(Relat_table_t_arm_t_cl).all()
            collums = ['id', 'arm_t_id', 'class_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "ID типа брони", "ID класса"])
        if result == "Связи заклинаний и классов":
            table = Relat_table_t_spell_cl
            line = 28  # session.query(Relat_table_t_spell_cl).count()
            self.ui.listView.setColumnCount(3)
            self.ui.qTable = session.query(Relat_table_t_spell_cl).all()
            collums = ['id', 'spell_id', 'class_id']
            self.ui.listView.setHorizontalHeaderLabels(["ID", "ID заклинания", "ID класса"])
        if result == "Персонаж":
            table = Character
            line = session.query(Character).count()
            self.ui.listView.setColumnCount(13)
            self.ui.qTable = session.query(Character).all()
            collums = ['id', 'name', 'power', 'agility', 'body_type', 'intellect', 'wisdom', 'charisma', 'acc_id',
                       'var_races_id', 'class_id', 'weap_id', 'arm_id']
            self.ui.listView.setHorizontalHeaderLabels(
                ["ID", "Имя", "Сила", "Ловкость", "Телосложение", "Интеллект", "Мудрость", "Харизма", "ID аккаунта",
                 "ID разновидности расы", "ID класса", "ID оружия", "ID брони"])

        self.ui.listView.setRowCount(line)
        sort = [0] * line
        j = 0
        for i in collums:
            if i != 'id':
                self.ui.comboBox_2.addItem(i)
        for i in session.query(table).all():
            sort[j] = i.id
            j += 1
        sort.sort()
        for i in sort:
            self.ui.comboBox_3.addItem(str(i))
        perem = 0
        j = -1
        for i in sort:
            j += 1
            for row, form in enumerate(self.ui.qTable):
                col = 0
                for c in collums:
                    for k, v in vars(form).items():
                        if c == k:
                            if c == 'id':
                                perem = v
                            if ((c == 'id') & (v == i)) | ((c != 'id') & (perem == i)):
                                self.ui.listView.setItem(j, col, QTableWidgetItem(str(v)))
                                col += 1
        self.ui.listView.setStyleSheet("selection-color: rgb(255, 0, 127);\n"
                                       "selection-background-color: rgb(85, 255, 127);")
        self.ui.listView.resizeColumnsToContents()

    # Изменение таблицы
    def change_table(self):
        self.ui.statusbar.clearMessage()
        if self.ui.prov == 0:
            self.ui.prov = 1
        result = self.ui.comboBox.currentText()
        change = self.ui.comboBox_4.currentText()
        stri = int(self.ui.comboBox_3.currentText())
        stlb = self.ui.comboBox_2.currentText()
        text = self.ui.textBrowser.toPlainText()
        lis = text.split(',')
        table = Weapon
        user_setting = insert(table).values(id=session.query(table).count() + 1, name=text)
        query = session.query(table).filter(table.id == 1)
        try:
            if result == "Типы оружия":
                table = weapon_types
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=text)
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    query.name = text
            if result == "Типы брони":
                table = armor_types
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=text)
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    query.name = text
            if result == "Статусы пользователя":
                table = user_status
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, status=text)
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    query.status = text
            if result == "Описания":
                table = descriptions
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, field=text)
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    query.field = text
            if result == "Расы":
                table = Races
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], incr_char=lis[1],
                                         worldview=lis[2], size=lis[3], speed=int(lis[4]), descr_id=int(lis[5]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "incr_char":
                        query.incr_char = text
                    if stlb == "worldview":
                        query.worldview = text
                    if stlb == "size":
                        query.size = text
                    if stlb == "speed":
                        query.speed = int(text)
                    if stlb == "descr_id":
                        query.descr_id = int(text)
            if result == "Разновидности рас":
                table = var_races
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], incr_char=lis[1],
                                         add_feat=lis[2], rac_id=int(lis[3]), descr_id=int(lis[4]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                        print(query.name)
                    if stlb == "incr_char":
                        query.incr_char = text
                    if stlb == "add_feat":
                        query.add_feat = text
                    if stlb == "rac_id":
                        query.rac_id = text
                    if stlb == "descr_id":
                        query.descr_id = int(text)
            if result == "Аккаунты":
                table = Accounts
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, login=lis[0], password=lis[1],
                                         stat_id=int(lis[2]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "login":
                        query.login = text
                    if stlb == "password":
                        query.password = text
                    if stlb == "stat_id":
                        query.stat_id = int(text)
            if result == "Классы":
                table = Classes
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0],
                                         master_bonus=int(lis[1]),
                                         numb_av_spells=int(lis[2]), descr_id=int(lis[3]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "master_bonus":
                        query.master_bonus = int(text)
                    if stlb == "numb_av_spells":
                        query.numb_av_spells = int(text)
                    if stlb == "descr_id":
                        query.descr_id = int(text)
            if result == "Оружие":
                table = Weapon
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], price=int(lis[1]),
                                         damage=int(lis[2]), weight=int(lis[3]), weap_t_id=int(lis[4]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "price":
                        query.price = int(text)
                    if stlb == "damage":
                        query.damage = int(text)
                    if stlb == "weight":
                        query.weight = int(text)
                    if stlb == "weap_t_id":
                        query.weap_t_id = int(text)
            if result == "Броня":
                table = Armor
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], price=int(lis[1]),
                                         steal_hindr=bool(lis[2]), weight=int(lis[3]), arm_t_id=int(lis[4]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "price":
                        query.price = int(text)
                    if stlb == "steal_hindr":
                        query.steal_hindr = bool(text)
                    if stlb == "weight":
                        query.weight = int(text)
                    if stlb == "arm_t_id":
                        query.arm_t_id = int(text)
            if result == "Таблица заклинаний":
                table = Spell_table
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], level=int(lis[1]),
                                         distance=int(lis[2]), duration=int(lis[3]), descr_id=int(lis[4]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "level":
                        query.level = int(text)
                    if stlb == "distance":
                        query.distance = int(text)
                    if stlb == "duration":
                        query.duration = int(text)
                    if stlb == "descr_id":
                        query.descr_id = int(text)
            if result == "Связи типов оружия и классов":
                table = Relat_table_t_weap_t_cl
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, weap_t_id=int(lis[0]),
                                         class_id=int(lis[1]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "weap_t_id":
                        query.weap_t_id = int(text)
                    if stlb == "class_id":
                        query.class_id = int(text)
            if result == "Связи типов доспехов и классов":
                table = Relat_table_t_arm_t_cl
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, arm_t_id=int(lis[0]),
                                         class_id=int(lis[1]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "arm_t_id":
                        query.arm_t_id = int(text)
                    if stlb == "class_id":
                        query.class_id = int(text)
            if result == "Связи заклинаний и классов":
                table = Relat_table_t_spell_cl
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, spell_id=int(lis[0]),
                                         class_id=int(lis[1]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "spell_id":
                        query.spell_id = int(text)
                    if stlb == "class_id":
                        query.class_id = int(text)
            if result == "Персонаж":
                table = Character
                if change == "Добавить строку":
                    f = 0  # флаг
                    line = session.query(table).count()
                    tabl = session.query(table).all()
                    free_id2 = 1  # свободный индекс
                    while f == 0:
                        for i in range(line):
                            if free_id2 == tabl[i].id:
                                f = 1
                        if f == 1:
                            free_id2 = free_id2 + 1
                            f = 0
                        else:
                            f = 1
                    user_setting = table(id=free_id2, name=lis[0], power=int(lis[1]),
                                         agility=int(lis[2]), body_type=int(lis[3]), intellect=int(lis[4]),
                                         wisdom=int(lis[5]), charisma=int(lis[6]), acc_id=int(lis[7]),
                                         var_races_id=int(lis[8]), class_id=int(lis[9]), weap_id=int(lis[10]),
                                         arm_id=int(lis[11]))
                if change == "Обновить элемент":
                    query = session.query(table).filter(table.id == stri).first()
                    if stlb == "name":
                        query.name = text
                    if stlb == "power":
                        query.power = int(text)
                    if stlb == "agility":
                        query.agility = int(text)
                    if stlb == "body_type":
                        query.body_type = int(text)
                    if stlb == "intellect":
                        query.intellect = int(text)
                    if stlb == "wisdom":
                        query.wisdom = int(text)
                    if stlb == "charisma":
                        query.charisma = int(text)
                    if stlb == "acc_id":
                        query.acc_id = int(text)
                    if stlb == "var_races_id":
                        query.var_races_id = int(text)
                    if stlb == "class_id":
                        query.class_id = int(text)
                    if stlb == "weap_id":
                        query.weap_id = int(text)
                    if stlb == "arm_id":
                        query.arm_id = int(text)
            if change == "Удалить строку":
                session.query(table).filter_by(id=stri).delete(synchronize_session=False)
            if change == "Добавить строку":
                session.add(user_setting)

            self.show_table()
        except:
            self.ui.statusbar.showMessage("Произошла ошибка")

    # Подтвердить изменения
    def confirm_change(self):
        if self.ui.prov != 0:
            session.commit()
            DumpPostgreSql()
            self.ui.prov = 0
            self.show_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    # Отменить изменения
    def undo_change(self):
        if self.ui.prov != 0:
            session.rollback()
            self.ui.prov = 0
            self.show_table()
        else:
            self.ui.statusbar.showMessage("Сначала внесите изменение в таблицу")

    # Выход из панели администратора обратно в меню
    def exit_db_panel(self):
        self.hide()
        dialog = AdminMenu(parent=self)
        dialog.show()


# Панель входа
class log_panel(QMainWindow):
    def __init__(self, parent=None):
        super(log_panel, self).__init__(parent)
        self.ui = login_panell()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda: self.inter_to_app())
        self.ui.pushButton_2.clicked.connect(lambda: self.create_acc())

    # Войти в аккаунт
    def inter_to_app(self):
        log = self.ui.textEdit.toPlainText()
        pasw = self.ui.textEdit_2.toPlainText()
        if log == "":
            self.ui.statusbar.showMessage("Введите логин")
            return
        if pasw == "":
            self.ui.statusbar.showMessage("Введите пароль")
            return
        for i in session.query(Accounts).all():
            if log == i.login:
                if pasw == i.password:
                    self.hide()
                    if i.stat_id == 1:
                        MyGlobals.id = i.id
                        dialog = AdminMenu(parent=self)
                        dialog.show()
                    else:
                        MyGlobals.id = i.id
                        dialog = PolsMenu(parent=self)
                        dialog.show()
                else:
                    self.ui.statusbar.showMessage("Введён неверный пароль")
                    return
        self.ui.statusbar.showMessage("Введено несуществующее имя пользователя")

    # Создать аккаунт
    def create_acc(self):
        log = self.ui.textEdit.toPlainText()
        pasw = self.ui.textEdit_2.toPlainText()
        if log == "":
            self.ui.statusbar.showMessage("Введите логин для нового пользователя")
            return
        if pasw == "":
            self.ui.statusbar.showMessage("Введите пароль для нового пользователя")
            return
        for i in session.query(Accounts).all():
            if log == i.login:
                self.ui.statusbar.showMessage("Пользователь с таким логином уже существует")
                return
        f = 0  # флаг
        line = session.query(Accounts).count()
        tabl = session.query(Accounts).all()
        free_id2 = 1  # свободный индекс
        while f == 0:
            for i in range(line):
                if free_id2 == tabl[i].id:
                    f = 1
            if f == 1:
                free_id2 = free_id2 + 1
                f = 0
            else:
                f = 1
        user_setting = Accounts(id=free_id2, login=log, password=pasw,
                                stat_id=int(2))
        MyGlobals.id = free_id2
        session.add(user_setting)
        session.commit()
        DumpPostgreSql()
        self.hide()
        dialog = PolsMenu(parent=self)
        dialog.show()


# Окно описания
class descri(QMainWindow):
    def __init__(self, parent=None):
        super(descri, self).__init__(parent)
        self.ui = descrip()
        self.ui.setupUi(self)

    def des_show(self, lol):
        self.ui.textBrowser.clear()
        new = session.query(descriptions).filter_by(id=lol.descr_id).first()
        self.ui.textBrowser.insertPlainText(new.field)

    def word_show(self, lol):
        self.ui.textBrowser.clear()
        self.ui.textBrowser.insertPlainText(lol.worldview)

    def size_show(self, lol):
        self.ui.textBrowser.clear()
        self.ui.textBrowser.insertPlainText(lol.size + "\n" + "Скорость: " + str(lol.speed) + " футов")

    def dop_sp_show(self, lol):
        self.ui.textBrowser.clear()
        self.ui.textBrowser.insertPlainText(lol.add_feat)


# Окно описания заклинания
class desc_spell(QMainWindow):
    def __init__(self, parent=None):
        super(desc_spell, self).__init__(parent)
        self.ui = descript_spell()
        self.ui.setupUi(self)

    def des_show(self, lol):
        spell = session.query(Spell_table).filter_by(id=lol).first()
        self.ui.textBrowser.clear()
        new = session.query(descriptions).filter_by(id=spell.descr_id).first()
        self.ui.textBrowser.insertPlainText(new.field)
        if spell.level == 1:
            self.ui.textBrowser_2.insertPlainText("Первый уровень")
        else:
            self.ui.textBrowser_2.insertPlainText("Заговор")
        self.ui.textBrowser_3.insertPlainText(str(spell.distance) + " футов")
        self.ui.textBrowser_4.insertPlainText(str(spell.duration) + " мин.")


# Окно создания персонажа
class CreatePers(QMainWindow):
    def __init__(self, parent=None):
        super(CreatePers, self).__init__(parent)
        self.ui = create_per()
        self.ui.setupUi(self)
        self.add_combo()
        self.ui.action.triggered.connect(lambda: self.ret_to_menu())
        self.ui.pushButton_5.clicked.connect(lambda: self.show_class_inform())
        self.ui.pushButton.clicked.connect(lambda: self.stand_ch())
        self.ui.pushButton_2.clicked.connect(lambda: self.rand_ch())
        self.ui.pushButton_6.clicked.connect(lambda: self.show_race_inform())
        self.ui.pushButton_7.clicked.connect(lambda: self.show_var_race_inform())
        self.ui.pushButton_9.clicked.connect(lambda: self.spell_inform())
        self.ui.pushButton_4.clicked.connect(lambda: self.show_arm_inform())
        self.ui.pushButton_3.clicked.connect(lambda: self.show_weap_inform())
        self.ui.pushButton_8.clicked.connect(lambda: self.create_pers())

    # Создание персонажа
    def create_pers(self):
        if self.ui.comboBox_5.currentText() == "":
            self.ui.statusbar.showMessage("Сначала выберите расу и её разновидность")
            return
        if self.ui.textEdit.toPlainText() == "":
            self.ui.statusbar.showMessage("Сначала введите имя персонажа")
            return
        if self.ui.textBrowser_9.toPlainText() == "":
            self.ui.statusbar.showMessage("Сначала выберите стандартные характеристики")
            return
        if self.ui.comboBox_2.currentText() == "":
            self.ui.statusbar.showMessage("Сначала выберите броню")
            return
        if self.ui.comboBox.currentText() == "":
            self.ui.statusbar.showMessage("Сначала выберите оружие")
            return
        named = self.ui.textEdit.toPlainText()
        if session.query(Character).filter_by(name=named, acc_id=MyGlobals.id).first():
            self.ui.statusbar.showMessage("Уже есть персонаж с таким именем")
            return
        arm_name = self.ui.comboBox_2.currentText()
        weap_name = self.ui.comboBox.currentText()
        class_name = self.ui.comboBox_3.currentText()
        var_race_name = self.ui.comboBox_5.currentText()
        new = session.query(var_races).filter_by(name=var_race_name).first()
        new_2 = session.query(Classes).filter_by(name=class_name).first()
        new_3 = session.query(Weapon).filter_by(name=weap_name).first()
        new_4 = session.query(Armor).filter_by(name=arm_name).first()
        f = 0  # флаг
        line = session.query(Character).count()
        tabl = session.query(Character).all()
        free_id2 = 1  # свободный индекс
        while f == 0:
            for i in range(line):
                if free_id2 == tabl[i].id:
                    f = 1
            if f == 1:
                free_id2 = free_id2 + 1
                f = 0
            else:
                f = 1
        user_setting = Character(id=free_id2, name=named,
                                 power=int(self.ui.textBrowser_9.toPlainText()),
                                 agility=int(self.ui.textBrowser_8.toPlainText()),
                                 body_type=int(self.ui.textBrowser_7.toPlainText()),
                                 intellect=int(self.ui.textBrowser_6.toPlainText()),
                                 wisdom=int(self.ui.textBrowser_5.toPlainText()),
                                 charisma=int(self.ui.textBrowser_4.toPlainText()), acc_id=MyGlobals.id,
                                 var_races_id=int(new.id), class_id=int(new_2.id), weap_id=int(new_3.id),
                                 arm_id=int(new_4.id))
        session.add(user_setting)
        session.commit()
        DumpPostgreSql()
        self.ret_to_menu()

    # Рандомные стартовые характеристики
    def rand_ch(self):
        self.ui.textBrowser_9.clear()
        self.ui.textBrowser_8.clear()
        self.ui.textBrowser_7.clear()
        self.ui.textBrowser_5.clear()
        self.ui.textBrowser_4.clear()
        self.ui.textBrowser_6.clear()
        for j in range(6):
            mas = [0] * 5
            mas[0] = random.randint(1, 6)
            mas[1] = random.randint(1, 6)
            mas[2] = random.randint(1, 6)
            mas[3] = random.randint(1, 6)
            mas.sort()
            if j == 0:
                self.ui.textBrowser_9.insertPlainText(str(mas[1] + mas[2] + mas[3]))
            if j == 1:
                self.ui.textBrowser_8.insertPlainText(str(mas[1] + mas[2] + mas[3]))
            if j == 2:
                self.ui.textBrowser_7.insertPlainText(str(mas[1] + mas[2] + mas[3]))
            if j == 3:
                self.ui.textBrowser_5.insertPlainText(str(mas[1] + mas[2] + mas[3]))
            if j == 4:
                self.ui.textBrowser_4.insertPlainText(str(mas[1] + mas[2] + mas[3]))
            if j == 5:
                self.ui.textBrowser_6.insertPlainText(str(mas[1] + mas[2] + mas[3]))

    # Стандартные стартовые характеристики
    def stand_ch(self):
        self.ui.textBrowser_9.clear()
        self.ui.textBrowser_8.clear()
        self.ui.textBrowser_7.clear()
        self.ui.textBrowser_5.clear()
        self.ui.textBrowser_4.clear()
        self.ui.textBrowser_6.clear()
        id_1 = random.randint(1, 3)
        if id_1 == 1:
            self.ui.textBrowser_9.insertPlainText("15")
            self.ui.textBrowser_8.insertPlainText("14")
            self.ui.textBrowser_7.insertPlainText("13")
            self.ui.textBrowser_5.insertPlainText("12")
            self.ui.textBrowser_4.insertPlainText("10")
            self.ui.textBrowser_6.insertPlainText("8")
        if id_1 == 2:
            self.ui.textBrowser_9.insertPlainText("10")
            self.ui.textBrowser_8.insertPlainText("14")
            self.ui.textBrowser_7.insertPlainText("8")
            self.ui.textBrowser_5.insertPlainText("12")
            self.ui.textBrowser_4.insertPlainText("13")
            self.ui.textBrowser_6.insertPlainText("15")
        if id_1 == 3:
            self.ui.textBrowser_9.insertPlainText("14")
            self.ui.textBrowser_8.insertPlainText("15")
            self.ui.textBrowser_7.insertPlainText("10")
            self.ui.textBrowser_5.insertPlainText("8")
            self.ui.textBrowser_4.insertPlainText("13")
            self.ui.textBrowser_6.insertPlainText("12")

    # Показать всё, что связано с выбранной бронёй
    def show_arm_inform(self):
        self.ui.textBrowser_16.clear()
        self.ui.textBrowser_17.clear()
        self.ui.textBrowser_18.clear()
        if self.ui.comboBox_2.currentText() == "":
            self.ui.statusbar.showMessage("Сначала выберите класс и выведите информацию о нём")
            return
        pop = session.query(Armor).filter_by(name=self.ui.comboBox_2.currentText()).first()
        self.ui.textBrowser_17.insertPlainText(str(pop.price) + " зм.")
        self.ui.textBrowser_16.insertPlainText(str(pop.weight) + " фнт.")
        if pop.steal_hindr:
            self.ui.textBrowser_18.insertPlainText("Присутствует")
        else:
            self.ui.textBrowser_18.insertPlainText("Отсутствует")

    # Показать всё, что связано с оружием
    def show_weap_inform(self):
        self.ui.textBrowser_11.clear()
        self.ui.textBrowser_12.clear()
        self.ui.textBrowser_14.clear()
        if self.ui.comboBox.currentText() == "":
            self.ui.statusbar.showMessage("Сначала выберите класс и выведите информацию о нём")
            return
        pop = session.query(Weapon).filter_by(name=self.ui.comboBox.currentText()).first()
        self.ui.textBrowser_12.insertPlainText(str(pop.price) + " зм.")
        self.ui.textBrowser_14.insertPlainText(str(pop.weight) + " фнт.")
        self.ui.textBrowser_11.insertPlainText(str(pop.damage))

    # Показать информацию о заклинании
    def spell_inform(self):
        spell = self.ui.comboBox_6.currentText()
        if spell == "":
            self.ui.statusbar.showMessage("У этого класса нет ни одного заклинания")
            return
        else:
            pop = session.query(Spell_table).filter_by(name=spell).first()
            dialog = desc_spell(parent=self)
            dialog.show()
            dialog.des_show(pop.id)

    # Добавить в комбобоксы элементы
    def add_combo(self):
        for i in session.query(Classes).all():
            self.ui.comboBox_3.addItem(i.name)
        for i in session.query(Races).all():
            self.ui.comboBox_4.addItem(i.name)

    # Вернуться обратно в меню
    def ret_to_menu(self):
        self.hide()
        pop = session.query(Accounts).filter_by(id=MyGlobals.id).first()
        if pop.stat_id == 1:
            dialog = AdminMenu(parent=self)
        else:
            dialog = PolsMenu(parent=self)
        dialog.show()

    # Показать всё, что связано с выбранной расой
    def show_race_inform(self):
        self.ui.textBrowser.clear()
        self.ui.comboBox_5.clear()
        new = session.query(Races).filter_by(name=self.ui.comboBox_4.currentText()).first()
        for i in session.query(var_races).all():
            if i.rac_id == new.id:
                self.ui.comboBox_5.addItem(i.name)
        new_2 = session.query(descriptions).filter_by(id=new.descr_id).first()
        self.ui.textBrowser.insertPlainText(new_2.field + "\n\nМировозрение:\n" + new.worldview + "\n\nРазмер:\n" +
                                            new.size + "\n\nСкорость: " + str(new.speed) + " футов" +
                                            "\n\nПовышение характеристик:\n" + new.incr_char)

    # Показать всё, что связано с выбранной разновидностью расы
    def show_var_race_inform(self):
        self.ui.textBrowser_3.clear()
        named = self.ui.comboBox_5.currentText()
        if named == "":
            self.ui.statusbar.showMessage("Сначала выведите информацию о расе")
            return
        new = session.query(var_races).filter_by(name=named).first()
        new_2 = session.query(descriptions).filter_by(id=new.descr_id).first()
        self.ui.textBrowser_3.insertPlainText(new_2.field + "\n\nДоп. способность:\n" + new.add_feat +
                                              "\n\nПовышение характеристик:\n" + new.incr_char)

    # Показать всё, что связано с выбранным классом
    def show_class_inform(self):
        self.ui.textBrowser_2.clear()
        self.ui.comboBox.clear()
        self.ui.comboBox_2.clear()
        self.ui.comboBox_6.clear()
        new = session.query(Classes).filter_by(name=self.ui.comboBox_3.currentText()).first()
        for i in session.query(Relat_table_t_arm_t_cl).all():
            if i.class_id == new.id:
                new_2 = session.query(armor_types).filter_by(id=i.arm_t_id).first()
                for j in session.query(Armor).all():
                    if j.arm_t_id == new_2.id:
                        self.ui.comboBox_2.addItem(j.name)
        for i in session.query(Relat_table_t_weap_t_cl).all():
            if i.class_id == new.id:
                new_2 = session.query(weapon_types).filter_by(id=i.weap_t_id).first()
                for j in session.query(Weapon).all():
                    if j.weap_t_id == new_2.id:
                        self.ui.comboBox.addItem(j.name)
        for i in session.query(Relat_table_t_spell_cl).all():
            if i.class_id == new.id:
                new_2 = session.query(Spell_table).filter_by(id=i.spell_id).first()
                self.ui.comboBox_6.addItem(new_2.name)
        new_2 = session.query(descriptions).filter_by(id=new.descr_id).first()
        self.ui.textBrowser_2.insertPlainText(new_2.field + "\n\nБонус мастерства: " + str(new.master_bonus))


# Главное меню администратора БД
class AdminMenu(QMainWindow):
    def __init__(self, parent=None):
        super(AdminMenu, self).__init__(parent)
        self.ui = adm_menu()
        self.ui.setupUi(self)
        self.dost_pers()
        self.ui.comboBox_4.addItem("О классе")
        self.ui.comboBox_4.addItem("О расе")
        self.ui.comboBox_4.addItem("О мировозрении")
        self.ui.comboBox_4.addItem("О размере и скорости")
        self.ui.comboBox_4.addItem("О разновидности расы")
        self.ui.comboBox_4.addItem("О дополнительной способности")
        self.ui.comboBox_4.addItem("О выбранном заклинании")
        self.ui.comboBox_2.addItem("В Docx")
        self.ui.comboBox_2.addItem("В Xlsx")
        self.ui.pushButton.clicked.connect(lambda: self.show_pers_inform())
        self.ui.pushButton_5.clicked.connect(lambda: self.go_to_db_pan())
        self.ui.pushButton_2.clicked.connect(lambda: self.create_p())
        self.ui.action.triggered.connect(lambda: self.out_acc())
        self.ui.pushButton_3.clicked.connect(lambda: self.dell_pers())
        self.ui.pushButton_6.clicked.connect(lambda: self.descr())
        self.ui.pushButton_4.clicked.connect(lambda: self.vugr_db())

    # Выгрузка БД в Docx или Xlsx
    def vugr_db(self):
        if self.ui.comboBox_2.currentText() == "В Docx":
            mydoc = docx.Document()
            for i in range(15):
                table_1 = mydoc.add_table(rows=0, cols=0, style='Table Grid')
                table = Weapon
                line = 0
                collums = []
                if i == 0:
                    mydoc.add_paragraph("Таблица типов оружия (weapon_types):")
                    table_1 = mydoc.add_table(rows=session.query(weapon_types).count() + 1, cols=2, style='Table Grid')
                    table = weapon_types
                    line = session.query(weapon_types).count()
                    self.ui.qTable = session.query(weapon_types).all()
                    collums = ['id', 'name']
                if i == 1:
                    mydoc.add_paragraph("Таблица типов брони (armor_types):")
                    table_1 = mydoc.add_table(rows=session.query(armor_types).count() + 1, cols=2, style='Table Grid')
                    table = armor_types
                    line = session.query(armor_types).count()
                    self.ui.qTable = session.query(armor_types).all()
                    collums = ['id', 'name']
                if i == 2:
                    mydoc.add_paragraph("Таблица статусов пользователя (user_status):")
                    table_1 = mydoc.add_table(rows=session.query(user_status).count() + 1, cols=2, style='Table Grid')
                    table = user_status
                    line = session.query(user_status).count()
                    self.ui.qTable = session.query(user_status).all()
                    collums = ['id', 'status']
                if i == 3:
                    mydoc.add_paragraph("Таблица описаний (descriptions):")
                    table_1 = mydoc.add_table(rows=session.query(descriptions).count() + 1, cols=2, style='Table Grid')
                    table = descriptions
                    line = session.query(descriptions).count()
                    self.ui.qTable = session.query(descriptions).all()
                    collums = ['id', 'field']
                if i == 4:
                    mydoc.add_paragraph("Таблица рас (Races):")
                    table_1 = mydoc.add_table(rows=session.query(Races).count() + 1, cols=7, style='Table Grid')
                    table = Races
                    line = session.query(Races).count()
                    self.ui.qTable = session.query(Races).all()
                    collums = ['id', 'name', 'incr_char', 'worldview', 'size', 'speed', 'descr_id']
                if i == 5:
                    mydoc.add_paragraph("Таблица разновидностей рас (var_races):")
                    table_1 = mydoc.add_table(rows=14 + 1, cols=6, style='Table Grid')
                    table = var_races
                    line = 14  # session.query(var_races).count()
                    self.ui.qTable = session.query(var_races).all()
                    collums = ['id', 'name', 'incr_char', 'add_feat', 'rac_id', 'descr_id']
                if i == 6:
                    mydoc.add_paragraph("Таблица аккаунтов (Accounts):")
                    table_1 = mydoc.add_table(rows=session.query(Accounts).count() + 1, cols=4, style='Table Grid')
                    table = Accounts
                    line = session.query(Accounts).count()
                    self.ui.qTable = session.query(Accounts).all()
                    collums = ['id', 'login', 'password', 'stat_id']
                if i == 7:
                    mydoc.add_paragraph("Таблица классов (Classes):")
                    table_1 = mydoc.add_table(rows=session.query(Classes).count() + 1, cols=5, style='Table Grid')
                    table = Classes
                    line = session.query(Classes).count()
                    self.ui.qTable = session.query(Classes).all()
                    collums = ['id', 'name', 'master_bonus', 'numb_av_spells', 'descr_id']
                if i == 8:
                    mydoc.add_paragraph("Таблица оружия (Classes):")
                    table_1 = mydoc.add_table(rows=session.query(Weapon).count() + 1, cols=6, style='Table Grid')
                    table = Weapon
                    line = session.query(Weapon).count()
                    self.ui.qTable = session.query(Weapon).all()
                    collums = ['id', 'name', 'price', 'damage', 'weight', 'weap_t_id']
                if i == 9:
                    mydoc.add_paragraph("Таблица брони (Armor):")
                    table_1 = mydoc.add_table(rows=session.query(Armor).count() + 1, cols=6, style='Table Grid')
                    table = Armor
                    line = session.query(Armor).count()
                    self.ui.qTable = session.query(Armor).all()
                    collums = ['id', 'name', 'price', 'steal_hindr', 'weight', 'arm_t_id']
                if i == 10:
                    mydoc.add_paragraph("Таблица заклинаний (Spell_table):")
                    table_1 = mydoc.add_table(rows=session.query(Spell_table).count() + 1, cols=6, style='Table Grid')
                    table = Spell_table
                    line = session.query(Spell_table).count()
                    self.ui.qTable = session.query(Spell_table).all()
                    collums = ['id', 'name', 'level', 'distance', 'duration', 'descr_id']
                if i == 11:
                    mydoc.add_paragraph("Таблица связи типов оружия и классов (Relat_table_t_weap_t_cl):")
                    table_1 = mydoc.add_table(rows=21 + 1, cols=3, style='Table Grid')
                    table = Relat_table_t_weap_t_cl
                    line = 21  # session.query(Relat_table_t_weap_t_cl).count()
                    self.ui.qTable = session.query(Relat_table_t_weap_t_cl).all()
                    collums = ['id', 'weap_t_id', 'class_id']
                if i == 12:
                    mydoc.add_paragraph("Таблица связи типов доспехов и классов (Relat_table_t_arm_t_cl):")
                    table_1 = mydoc.add_table(rows=25 + 1, cols=3,
                                              style='Table Grid')
                    table = Relat_table_t_arm_t_cl
                    line = 25  # session.query(Relat_table_t_arm_t_cl).count()
                    self.ui.qTable = session.query(Relat_table_t_arm_t_cl).all()
                    collums = ['id', 'arm_t_id', 'class_id']
                if i == 13:
                    mydoc.add_paragraph("Таблица связи заклинаний и классов (Relat_table_t_spell_cl):")
                    table_1 = mydoc.add_table(rows=28 + 1, cols=3,
                                              style='Table Grid')
                    table = Relat_table_t_spell_cl
                    line = 28  # session.query(Relat_table_t_spell_cl).count()
                    self.ui.qTable = session.query(Relat_table_t_spell_cl).all()
                    collums = ['id', 'spell_id', 'class_id']
                if i == 14:
                    mydoc.add_paragraph("Таблица персонажей (Character):")
                    table_1 = mydoc.add_table(rows=session.query(Character).count() + 1, cols=13,
                                              style='Table Grid')
                    table = Character
                    line = session.query(Character).count()
                    self.ui.qTable = session.query(Character).all()
                    collums = ['id', 'name', 'power', 'agility', 'body_type', 'intellect', 'wisdom', 'charisma',
                               'acc_id', 'var_races_id', 'class_id', 'weap_id', 'arm_id']
                z = 0
                for j in collums:
                    table_1.cell(0, z).text = j
                    z += 1
                sort = [0] * line
                z = 0
                for j in session.query(table).all():
                    sort[z] = j.id
                    z += 1
                sort.sort()
                perem = 0
                z = 0
                for j in sort:
                    z += 1
                    for row, form in enumerate(self.ui.qTable):
                        col = 0
                        for c in collums:
                            for k, v in vars(form).items():
                                if c == k:
                                    if c == 'id':
                                        perem = v
                                    if ((c == 'id') & (v == j)) | ((c != 'id') & (perem == j)):
                                        table_1.cell(z, col).text = str(v)
                                        col += 1
            stri = "/Users/chronos_386/Desktop/database_"
            for i in range(15):
                stri += str(random.randint(0, 9))
            mydoc.save(stri + ".docx")
        else:
            table1 = pd.DataFrame({})
            table2 = pd.DataFrame({})
            table3 = pd.DataFrame({})
            table4 = pd.DataFrame({})
            table5 = pd.DataFrame({})
            table6 = pd.DataFrame({})
            table7 = pd.DataFrame({})
            table8 = pd.DataFrame({})
            table9 = pd.DataFrame({})
            table10 = pd.DataFrame({})
            table11 = pd.DataFrame({})
            table12 = pd.DataFrame({})
            table13 = pd.DataFrame({})
            table14 = pd.DataFrame({})
            table15 = pd.DataFrame({})
            for i in range(15):
                if i == 0:
                    count = session.query(weapon_types).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(weapon_types).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        z += 1
                    table1 = pd.DataFrame({'id': collum1, 'Name': collum2})
                if i == 1:
                    count = session.query(armor_types).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(armor_types).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        z += 1
                    table2 = pd.DataFrame({'id': collum1, 'Name': collum2})
                if i == 2:
                    count = session.query(user_status).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(user_status).all():
                        collum1[z] = j.id
                        collum2[z] = j.status
                        z += 1
                    table3 = pd.DataFrame({'id': collum1, 'status': collum2})
                if i == 3:
                    count = session.query(descriptions).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    z = 0
                    for j in session.query(descriptions).all():
                        collum1[z] = j.id
                        collum2[z] = j.field
                        z += 1
                    table4 = pd.DataFrame({'id': collum1, 'field': collum2})
                if i == 4:
                    count = session.query(Races).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [""] * count
                    collum5 = [""] * count
                    collum6 = [0] * count
                    collum7 = [0] * count
                    z = 0
                    for j in session.query(Races).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.incr_char
                        collum4[z] = j.worldview
                        collum5[z] = j.size
                        collum6[z] = j.speed
                        collum7[z] = j.descr_id
                        z += 1
                    table5 = pd.DataFrame({'id': collum1, 'name': collum2, 'incr_char': collum3, 'worldview': collum4,
                                           'size': collum5, 'speed': collum6, 'descr_id': collum7})
                if i == 5:
                    count = 14  # session.query(var_races).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [""] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    z = 0
                    for j in session.query(var_races).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.incr_char
                        collum4[z] = j.add_feat
                        collum5[z] = j.rac_id
                        collum6[z] = j.descr_id
                        z += 1
                    table6 = pd.DataFrame({'id': collum1, 'name': collum2, 'incr_char': collum3, 'add_feat': collum4,
                                           'rac_id': collum5, 'descr_id': collum6})
                if i == 6:
                    count = session.query(Accounts).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [""] * count
                    collum4 = [0] * count
                    z = 0
                    for j in session.query(Accounts).all():
                        collum1[z] = j.id
                        collum2[z] = j.login
                        collum3[z] = j.password
                        collum4[z] = j.stat_id
                        z += 1
                    table7 = pd.DataFrame({'id': collum1, 'login': collum2, 'password': collum3, 'stat_id': collum4})
                if i == 7:
                    count = session.query(Classes).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    z = 0
                    for j in session.query(Classes).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.master_bonus
                        collum4[z] = j.numb_av_spells
                        collum5[z] = j.descr_id
                        z += 1
                    table8 = pd.DataFrame({'id': collum1, 'name': collum2, 'master_bonus': collum3,
                                           'numb_av_spells': collum4, 'descr_id': collum5})
                if i == 8:
                    count = session.query(Weapon).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    z = 0
                    for j in session.query(Weapon).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.price
                        collum4[z] = j.damage
                        collum5[z] = j.weight
                        collum6[z] = j.weap_t_id
                        z += 1
                    table9 = pd.DataFrame({'id': collum1, 'name': collum2, 'price': collum3,
                                           'damage': collum4, 'weight': collum5, 'weap_t_id': collum6})
                if i == 9:
                    count = session.query(Armor).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    z = 0
                    for j in session.query(Armor).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.price
                        collum4[z] = j.steal_hindr
                        collum5[z] = j.weight
                        collum6[z] = j.arm_t_id
                        z += 1
                    table10 = pd.DataFrame({'id': collum1, 'name': collum2, 'price': collum3,
                                            'steal_hindr': collum4, 'weight': collum5, 'arm_t_id': collum6})
                if i == 10:
                    count = session.query(Spell_table).count()
                    collum1 = [0] * count
                    collum2 = [""] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    z = 0
                    for j in session.query(Spell_table).all():
                        collum1[z] = j.id
                        collum2[z] = j.name
                        collum3[z] = j.level
                        collum4[z] = j.distance
                        collum5[z] = j.duration
                        collum6[z] = j.descr_id
                        z += 1
                    table11 = pd.DataFrame({'id': collum1, 'name': collum2, 'level': collum3,
                                            'distance': collum4, 'duration': collum5, 'descr_id': collum6})
                if i == 11:
                    count = 21  # session.query(Relat_table_t_weap_t_cl).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    z = 0
                    for j in session.query(Relat_table_t_weap_t_cl).all():
                        collum1[z] = j.id
                        collum2[z] = j.weap_t_id
                        collum3[z] = j.class_id
                        z += 1
                    table12 = pd.DataFrame({'id': collum1, 'weap_t_id': collum2, 'class_id': collum3})
                if i == 12:
                    count = 25  # session.query(Relat_table_t_arm_t_cl).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    z = 0
                    for j in session.query(Relat_table_t_arm_t_cl).all():
                        collum1[z] = j.id
                        collum2[z] = j.arm_t_id
                        collum3[z] = j.class_id
                        z += 1
                    table13 = pd.DataFrame({'id': collum1, 'arm_t_id': collum2, 'class_id': collum3})
                if i == 13:
                    count = 28  # session.query(Relat_table_t_spell_cl).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    z = 0
                    for j in session.query(Relat_table_t_spell_cl).all():
                        collum1[z] = j.id
                        collum2[z] = j.spell_id
                        collum3[z] = j.class_id
                        z += 1
                    table14 = pd.DataFrame({'id': collum1, 'spell_id': collum2, 'class_id': collum3})
                if i == 14:
                    count = session.query(Character).count()
                    collum1 = [0] * count
                    collum2 = [0] * count
                    collum3 = [0] * count
                    collum4 = [0] * count
                    collum5 = [0] * count
                    collum6 = [0] * count
                    collum7 = [0] * count
                    collum8 = [0] * count
                    collum9 = [0] * count
                    collum10 = [0] * count
                    collum11 = [0] * count
                    collum12 = [0] * count
                    collum13 = [0] * count
                    z = 0
                    for j in session.query(Character).all():
                        collum1[z] = j.id
                        collum2[z] = j.power
                        collum3[z] = j.agility
                        collum4[z] = j.body_type
                        collum5[z] = j.intellect
                        collum6[z] = j.wisdom
                        collum7[z] = j.charisma
                        collum8[z] = j.acc_id
                        collum9[z] = j.var_races_id
                        collum10[z] = j.class_id
                        collum11[z] = j.weap_id
                        collum12[z] = j.arm_id
                        collum13[z] = j.name
                        z += 1
                    table15 = pd.DataFrame({'id': collum1, 'name': collum13, 'power': collum2,
                                            'agility': collum3, 'body_type': collum4, 'intellect': collum5,
                                            'wisdom': collum6, 'charisma': collum7, 'acc_id': collum8,
                                            'var_races_id': collum9, 'class_id': collum10, 'weap_id': collum11,
                                            'arm_id': collum12})
            salary_sheets = {'Типы оружия': table1, 'Типы брони': table2, 'Статусы пользователей': table3,
                             'Описания': table4, 'Расы': table5, 'Разновидности рас': table6,
                             'Аккаунты': table7, 'Классы': table8, 'Оружие': table9, 'Броня': table10,
                             'Таблица заклинаний': table11, 'Отношения типов оржия и классов': table12,
                             'Отношения типов брони и классов': table13, 'Отношения заклинаний и классов': table14,
                             'Таблица персонажей': table15}
            writer = pd.ExcelWriter('./database123456.xlsx', engine='xlsxwriter')
            for sheet_name in salary_sheets.keys():
                salary_sheets[sheet_name].to_excel(writer, sheet_name=sheet_name, index=False)
            writer.save()

    # Переходим в окно создания персонажа
    def create_p(self):
        self.hide()
        dialog = CreatePers(parent=self)
        dialog.show()

    # Показ доступных персонажей
    def dost_pers(self):
        self.ui.comboBox.clear()
        for i in session.query(Character).all():
            if i.acc_id == MyGlobals.id:
                self.ui.comboBox.addItem(i.name)

    # Выход из аккаунта
    def out_acc(self):
        self.hide()
        dialog = log_panel(parent=self)
        dialog.show()

    # Переход в панель администратора БД
    def go_to_db_pan(self):
        self.hide()
        dialog = Admin_db_Panel(parent=self)
        dialog.show()

    # Удаление персонажа
    def dell_pers(self):
        named = self.ui.comboBox.currentText()
        pop = session.query(Character).filter_by(name=named, acc_id=MyGlobals.id).first()
        session.query(Character).filter_by(id=pop.id).delete(synchronize_session=False)
        session.commit()
        DumpPostgreSql()
        self.clear_all()
        self.ui.comboBox.clear()
        self.dost_pers()

    # Очистка всего окна
    def clear_all(self):
        self.ui.textBrowser_9.clear()
        self.ui.textBrowser_8.clear()
        self.ui.textBrowser_7.clear()
        self.ui.textBrowser_5.clear()
        self.ui.textBrowser_4.clear()
        self.ui.textBrowser_6.clear()
        self.ui.textBrowser.clear()
        self.ui.comboBox_3.clear()
        self.ui.textBrowser_3.clear()
        self.ui.textBrowser_2.clear()
        self.ui.textBrowser_10.clear()
        self.ui.textBrowser_11.clear()
        self.ui.textBrowser_12.clear()
        self.ui.textBrowser_14.clear()
        self.ui.textBrowser_13.clear()
        self.ui.textBrowser_15.clear()
        self.ui.textBrowser_17.clear()
        self.ui.textBrowser_18.clear()
        self.ui.textBrowser_16.clear()
        self.ui.textBrowser_19.clear()

    # Узнать о персонаже подробнее
    def descr(self):
        spell = self.ui.comboBox_3.currentText()
        opre = self.ui.comboBox_4.currentText()
        if self.ui.textBrowser_2.toPlainText() == "":
            self.ui.statusbar.showMessage("Сначала выведите информацию о персонаже")
            return
        if opre == "О выбранном заклинании":
            if spell == "":
                self.ui.statusbar.showMessage("У вас нет ни одного заклинания")
                return
            else:
                pop = session.query(Spell_table).filter_by(name=spell).first()
                dialog = desc_spell(parent=self)
                dialog.show()
                dialog.des_show(pop.id)
        zapom = session.query(Character).filter_by(acc_id=MyGlobals.id, name=self.ui.comboBox.currentText()).first()
        save = session.query(var_races).filter_by(id=zapom.var_races_id).first()
        save_2 = session.query(Races).filter_by(id=save.rac_id).first()
        if opre == "О классе":
            named = self.ui.textBrowser_2.toPlainText()
            pop = session.query(Classes).filter_by(name=named).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О расе":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О разновидности расы":
            pop = session.query(var_races).filter_by(name=save.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О мировозрении":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.word_show(pop)
        if opre == "О размере и скорости":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.size_show(pop)
        if opre == "О дополнительной способности":
            pop = session.query(var_races).filter_by(name=save.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.dop_sp_show(pop)

    # Показать информацио о персонаже
    def show_pers_inform(self):
        self.clear_all()
        named = self.ui.comboBox.currentText()
        if named == "":
            self.ui.statusbar.showMessage("У Вас ещё нет ни одного персонажа")
            return
        self.ui.textBrowser.insertPlainText(named)
        zapom = session.query(Character).filter_by(acc_id=MyGlobals.id, name=named).first()
        for i in session.query(Relat_table_t_spell_cl).all():
            if i.class_id == zapom.class_id:
                for j in session.query(Spell_table).all():
                    if j.id == i.spell_id:
                        self.ui.comboBox_3.addItem(j.name)
        save = session.query(var_races).filter_by(id=zapom.var_races_id).first()
        save_2 = session.query(Races).filter_by(id=save.rac_id).first()
        pov = save_2.incr_char.split('+')
        pove = 0
        agi = 0
        body = 0
        wisd = 0
        chr = 0
        intell = 0
        if pov[0] == "Сила":
            pove += int(pov[1])
        if pov[0] == "Ловкость":
            agi += int(pov[1])
        if pov[0] == "Телосложение":
            body += int(pov[1])
        if pov[0] == "Мудрость":
            wisd += int(pov[1])
        if pov[0] == "Харизма":
            chr += int(pov[1])
        if pov[0] == "Интеллект":
            intell += int(pov[1])
        if pov[0] == "Все":
            pove += int(pov[1])
            agi += int(pov[1])
            body += int(pov[1])
            wisd += int(pov[1])
            chr += int(pov[1])
            intell += int(pov[1])
        pov = save.incr_char.split('+')
        if pov[0] == "Сила":
            pove += int(pov[1])
        if pov[0] == "Ловкость":
            agi += int(pov[1])
        if pov[0] == "Телосложение":
            body += int(pov[1])
        if pov[0] == "Мудрость":
            wisd += int(pov[1])
        if pov[0] == "Харизма":
            chr += int(pov[1])
        if pov[0] == "Интеллект":
            intell += int(pov[1])
        if pov[0] == "Все":
            pove += int(pov[1])
            agi += int(pov[1])
            body += int(pov[1])
            wisd += int(pov[1])
            chr += int(pov[1])
            intell += int(pov[1])
        self.ui.textBrowser_9.insertPlainText(str(zapom.power + pove))
        self.ui.textBrowser_8.insertPlainText(str(zapom.agility + agi))
        self.ui.textBrowser_7.insertPlainText(str(zapom.body_type + body))
        self.ui.textBrowser_5.insertPlainText(str(zapom.wisdom + wisd))
        self.ui.textBrowser_4.insertPlainText(str(zapom.charisma + chr))
        self.ui.textBrowser_6.insertPlainText(str(zapom.intellect + intell))
        self.ui.textBrowser_3.insertPlainText(save_2.name + " " + save.name)
        save = session.query(Classes).filter_by(id=zapom.class_id).first()
        self.ui.textBrowser_2.insertPlainText(save.name)
        save = session.query(Weapon).filter_by(id=zapom.weap_id).first()
        self.ui.textBrowser_10.insertPlainText(save.name)
        self.ui.textBrowser_11.insertPlainText(str(save.damage))
        self.ui.textBrowser_12.insertPlainText(str(save.price) + " зм.")
        self.ui.textBrowser_14.insertPlainText(str(save.weight) + " фнт.")
        save_2 = session.query(weapon_types).filter_by(id=save.weap_t_id).first()
        self.ui.textBrowser_13.insertPlainText(save_2.name)
        save = session.query(Armor).filter_by(id=zapom.arm_id).first()
        self.ui.textBrowser_15.insertPlainText(save.name)
        self.ui.textBrowser_17.insertPlainText(str(save.price) + " зм.")
        if save.steal_hindr:
            self.ui.textBrowser_18.insertPlainText("Присутствует")
        else:
            self.ui.textBrowser_18.insertPlainText("Отсутствует")
        self.ui.textBrowser_16.insertPlainText(str(save.weight) + " фнт.")
        save_2 = session.query(armor_types).filter_by(id=save.arm_t_id).first()
        self.ui.textBrowser_19.insertPlainText(save_2.name)


# Меню обычного пользователя
class PolsMenu(QMainWindow):
    def __init__(self, parent=None):
        super(PolsMenu, self).__init__(parent)
        self.ui = pols_main_menu()
        self.ui.setupUi(self)
        self.dost_pers()
        self.ui.comboBox_4.addItem("О классе")
        self.ui.comboBox_4.addItem("О расе")
        self.ui.comboBox_4.addItem("О разновидности расы")
        self.ui.comboBox_4.addItem("О мировозрении")
        self.ui.comboBox_4.addItem("О дополнительной способности")
        self.ui.comboBox_4.addItem("О размере и скорости")
        self.ui.comboBox_4.addItem("О выбранном заклинании")
        self.ui.pushButton.clicked.connect(lambda: self.show_pers_inform())
        self.ui.action.triggered.connect(lambda: self.out_acc())
        self.ui.pushButton_3.clicked.connect(lambda: self.dell_pers())
        self.ui.pushButton_2.clicked.connect(lambda: self.create_p())
        self.ui.pushButton_4.clicked.connect(lambda: self.descr())

    # Переход в создание персонажа
    def create_p(self):
        self.hide()
        dialog = CreatePers(parent=self)
        dialog.show()

    # Описание чего-то конкретного о персонаже
    def descr(self):
        spell = self.ui.comboBox_3.currentText()
        opre = self.ui.comboBox_4.currentText()
        if self.ui.textBrowser_2.toPlainText() == "":
            self.ui.statusbar.showMessage("Сначала выведите информацию о персонаже")
            return
        if opre == "О выбранном заклинании":
            if spell == "":
                self.ui.statusbar.showMessage("У вас нет ни одного заклинания")
                return
            else:
                pop = session.query(Spell_table).filter_by(name=spell).first()
                dialog = desc_spell(parent=self)
                dialog.show()
                dialog.des_show(pop.id)
        zapom = session.query(Character).filter_by(acc_id=MyGlobals.id, name=self.ui.comboBox.currentText()).first()
        save = session.query(var_races).filter_by(id=zapom.var_races_id).first()
        save_2 = session.query(Races).filter_by(id=save.rac_id).first()
        if opre == "О классе":
            named = self.ui.textBrowser_2.toPlainText()
            pop = session.query(Classes).filter_by(name=named).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О расе":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О разновидности расы":
            pop = session.query(var_races).filter_by(name=save.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.des_show(pop)
        if opre == "О мировозрении":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.word_show(pop)
        if opre == "О размере и скорости":
            pop = session.query(Races).filter_by(name=save_2.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.size_show(pop)
        if opre == "О дополнительной способности":
            pop = session.query(var_races).filter_by(name=save.name).first()
            dialog = descri(parent=self)
            dialog.show()
            dialog.dop_sp_show(pop)

    # Показ доступных персонажей
    def dost_pers(self):
        self.ui.comboBox.clear()
        for i in session.query(Character).all():
            if i.acc_id == MyGlobals.id:
                self.ui.comboBox.addItem(i.name)

    # Очистка всех элементов окна
    def clear_all(self):
        self.ui.textBrowser_9.clear()
        self.ui.textBrowser_8.clear()
        self.ui.textBrowser_7.clear()
        self.ui.textBrowser_5.clear()
        self.ui.textBrowser_4.clear()
        self.ui.textBrowser_6.clear()
        self.ui.textBrowser.clear()
        self.ui.comboBox_3.clear()
        self.ui.textBrowser_3.clear()
        self.ui.textBrowser_2.clear()
        self.ui.textBrowser_10.clear()
        self.ui.textBrowser_11.clear()
        self.ui.textBrowser_12.clear()
        self.ui.textBrowser_14.clear()
        self.ui.textBrowser_13.clear()
        self.ui.textBrowser_15.clear()
        self.ui.textBrowser_17.clear()
        self.ui.textBrowser_18.clear()
        self.ui.textBrowser_16.clear()
        self.ui.textBrowser_19.clear()

    # Удаление персонажа
    def dell_pers(self):
        named = self.ui.comboBox.currentText()
        pop = session.query(Character).filter_by(name=named, acc_id=MyGlobals.id).first()
        session.query(Character).filter_by(id=pop.id).delete(synchronize_session=False)
        session.commit()
        DumpPostgreSql()
        self.clear_all()
        self.ui.comboBox.clear()
        self.dost_pers()

    # Выход из аккаунта
    def out_acc(self):
        self.hide()
        dialog = log_panel(parent=self)
        dialog.show()

    # Показ информации о персонаже
    def show_pers_inform(self):
        self.clear_all()
        named = self.ui.comboBox.currentText()
        if named == "":
            self.ui.statusbar.showMessage("У Вас ещё нет ни одного персонажа")
            return
        self.ui.textBrowser.insertPlainText(named)
        zapom = session.query(Character).filter_by(acc_id=MyGlobals.id, name=named).first()
        for i in session.query(Relat_table_t_spell_cl).all():
            if i.class_id == zapom.class_id:
                for j in session.query(Spell_table).all():
                    if j.id == i.spell_id:
                        self.ui.comboBox_3.addItem(j.name)
        save = session.query(var_races).filter_by(id=zapom.var_races_id).first()
        save_2 = session.query(Races).filter_by(id=save.rac_id).first()
        pov = save_2.incr_char.split('+')
        pove = 0
        agi = 0
        body = 0
        wisd = 0
        chr = 0
        intell = 0
        if pov[0] == "Сила":
            pove += int(pov[1])
        if pov[0] == "Ловкость":
            agi += int(pov[1])
        if pov[0] == "Телосложение":
            body += int(pov[1])
        if pov[0] == "Мудрость":
            wisd += int(pov[1])
        if pov[0] == "Харизма":
            chr += int(pov[1])
        if pov[0] == "Интеллект":
            intell += int(pov[1])
        if pov[0] == "Все":
            pove += int(pov[1])
            agi += int(pov[1])
            body += int(pov[1])
            wisd += int(pov[1])
            chr += int(pov[1])
            intell += int(pov[1])
        pov = save.incr_char.split('+')
        if pov[0] == "Сила":
            pove += int(pov[1])
        if pov[0] == "Ловкость":
            agi += int(pov[1])
        if pov[0] == "Телосложение":
            body += int(pov[1])
        if pov[0] == "Мудрость":
            wisd += int(pov[1])
        if pov[0] == "Харизма":
            chr += int(pov[1])
        if pov[0] == "Интеллект":
            intell += int(pov[1])
        if pov[0] == "Все":
            pove += int(pov[1])
            agi += int(pov[1])
            body += int(pov[1])
            wisd += int(pov[1])
            chr += int(pov[1])
            intell += int(pov[1])
        self.ui.textBrowser_9.insertPlainText(str(zapom.power + pove))
        self.ui.textBrowser_8.insertPlainText(str(zapom.agility + agi))
        self.ui.textBrowser_7.insertPlainText(str(zapom.body_type + body))
        self.ui.textBrowser_5.insertPlainText(str(zapom.wisdom + wisd))
        self.ui.textBrowser_4.insertPlainText(str(zapom.charisma + chr))
        self.ui.textBrowser_6.insertPlainText(str(zapom.intellect + intell))
        self.ui.textBrowser_3.insertPlainText(save_2.name + " " + save.name)
        save = session.query(Classes).filter_by(id=zapom.class_id).first()
        self.ui.textBrowser_2.insertPlainText(save.name)
        save = session.query(Weapon).filter_by(id=zapom.weap_id).first()
        self.ui.textBrowser_10.insertPlainText(save.name)
        self.ui.textBrowser_11.insertPlainText(str(save.damage))
        self.ui.textBrowser_12.insertPlainText(str(save.price) + " зм.")
        self.ui.textBrowser_14.insertPlainText(str(save.weight) + " фнт.")
        save_2 = session.query(weapon_types).filter_by(id=save.weap_t_id).first()
        self.ui.textBrowser_13.insertPlainText(save_2.name)
        save = session.query(Armor).filter_by(id=zapom.arm_id).first()
        self.ui.textBrowser_15.insertPlainText(save.name)
        self.ui.textBrowser_17.insertPlainText(str(save.price) + " зм.")
        if save.steal_hindr:
            self.ui.textBrowser_18.insertPlainText("Присутствует")
        else:
            self.ui.textBrowser_18.insertPlainText("Отсутствует")
        self.ui.textBrowser_16.insertPlainText(str(save.weight) + " фнт.")
        save_2 = session.query(armor_types).filter_by(id=save.arm_t_id).first()
        self.ui.textBrowser_19.insertPlainText(save_2.name)
